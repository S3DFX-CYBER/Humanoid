"""Render service — DOCX generation and PDF conversion.

Runs as an isolated container. Accepts structured document data and
produces formatted DOCX/PDF files.  No persistent state.
"""

import io
import logging
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import FastAPI
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

logger = logging.getLogger(__name__)

app = FastAPI(
    title="Humanoid Render Service",
    description="Document formatting — DOCX generation and PDF conversion",
    version="0.2.0",
)


class RenderRequest(BaseModel):
    """Input for document rendering."""

    job_id: str
    title: str = "Untitled Document"
    sections: list[dict] = []
    citation_style: str = "apa"
    references: list[dict] = []


def _build_docx(body: RenderRequest) -> Document:
    """Build a python-docx Document from structured data."""
    doc = Document()

    # -- Style defaults --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    para_format = style.paragraph_format
    para_format.space_after = Pt(6)
    para_format.line_spacing = 1.5

    # -- Title page --
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(120)
    run = title_para.add_run(body.title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = "Times New Roman"

    doc.add_page_break()

    # -- Sections --
    for section in body.sections:
        header = section.get("header", "")
        section_body = section.get("body", "")

        heading = doc.add_heading(header, level=2)
        heading.style.font.name = "Times New Roman"

        # Split into paragraphs on double-newline
        paragraphs = [p.strip() for p in section_body.split("\n\n") if p.strip()]
        for para_text in paragraphs:
            doc.add_paragraph(para_text)

    # -- References / Bibliography --
    if body.references:
        doc.add_page_break()
        doc.add_heading("References", level=1)

        for i, ref in enumerate(body.references, 1):
            title = ref.get("title", "Untitled")
            url = ref.get("url", "")
            ref_text = f"[{i}] {title}"
            if url:
                ref_text += f". Retrieved from {url}"
            doc.add_paragraph(ref_text)

    return doc


@app.get("/health")
async def health():
    return {"status": "ok", "service": "humanoid-render"}


@app.post("/render/docx")
async def render_docx(body: RenderRequest):
    """Generate a formatted DOCX from structured document data."""
    doc = _build_docx(body)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"humanoid_{body.job_id[:8]}.docx"

    return StreamingResponse(
        buffer,
        media_type=(
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/render/pdf")
async def render_pdf(body: RenderRequest):
    """Convert structured data to PDF via DOCX → LibreOffice."""
    doc = _build_docx(body)

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = Path(tmpdir) / "document.docx"
        doc.save(str(docx_path))

        # Run LibreOffice headless conversion
        result = subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                tmpdir,
                str(docx_path),
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )

        if result.returncode != 0:
            logger.error("soffice conversion failed: %s", result.stderr)
            return {
                "status": "error",
                "job_id": body.job_id,
                "message": f"PDF conversion failed: {result.stderr}",
            }

        pdf_path = Path(tmpdir) / "document.pdf"
        if not pdf_path.exists():
            return {
                "status": "error",
                "job_id": body.job_id,
                "message": "PDF file not found after conversion",
            }

        pdf_bytes = pdf_path.read_bytes()

    buffer = io.BytesIO(pdf_bytes)
    filename = f"humanoid_{body.job_id[:8]}.pdf"

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
